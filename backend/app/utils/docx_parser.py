import io
import os
import tempfile
from pathlib import Path
from typing import Any

from docx import Document
from PyPDF2 import PdfReader

try:
    import pypandoc
except ImportError:  # pragma: no cover
    pypandoc = None


def _to_paragraphs(text: str) -> list[str]:
    return [line.strip() for line in text.replace("\r", "").splitlines() if line.strip()]


def parse_docx_bytes(data: bytes) -> dict:
    stream = io.BytesIO(data)
    doc = Document(stream)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    tables = []
    for table in doc.tables:
        rows = []
        for r in table.rows:
            cells = [c.text for c in r.cells]
            rows.append(cells)
        tables.append(rows)
    return {"paragraphs": paragraphs, "tables": tables}


def parse_pdf_bytes(data: bytes) -> dict:
    reader = PdfReader(io.BytesIO(data))
    paragraphs: list[str] = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            paragraphs.extend(_to_paragraphs(text))
    return {"paragraphs": paragraphs, "tables": []}


def parse_doc_bytes(data: bytes) -> dict:
    if pypandoc is None:
        raise RuntimeError(
            "DOC parsing requires pypandoc and pandoc installed. "
            "Install pypandoc and the pandoc binary to parse .doc files."
        )
    with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as tmp:
        tmp.write(data)
        tmp.flush()
        tmp_path = tmp.name
    try:
        text = pypandoc.convert_file(tmp_path, "plain")
        paragraphs = _to_paragraphs(text)
        return {"paragraphs": paragraphs, "tables": []}
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass


def parse_resume_bytes(data: bytes, filename: str) -> dict:
    suffix = Path(filename).suffix.lower()
    if suffix == ".docx":
        return parse_docx_bytes(data)
    if suffix == ".pdf":
        return parse_pdf_bytes(data)
    if suffix == ".doc":
        return parse_doc_bytes(data)
    raise RuntimeError("Unsupported file type. Please upload .doc, .docx, or .pdf")
