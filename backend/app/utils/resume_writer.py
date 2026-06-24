import io
from typing import Any, Dict, List

from docx import Document


def _add_list(doc: Document, items: List[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _safe_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, list):
        return "; ".join(str(item) for item in value if item is not None)
    return str(value)


def resume_to_docx_bytes(pipeline: Dict[str, Any]) -> bytes:
    doc = Document()
    doc.add_heading("Enhanced Resume", level=0)

    profile = pipeline.get("profile", {}) or {}
    if profile:
        doc.add_heading("Profile Summary", level=1)
        for label, key in [
            ("Candidate Level", "candidate_level"),
            ("Primary Domain", "primary_domain"),
            ("Years of Experience", "years_experience"),
        ]:
            value = profile.get(key)
            if value:
                doc.add_paragraph(f"{label}: {_safe_text(value)}")

        skills = profile.get("skills")
        if skills:
            doc.add_heading("Core Skills", level=2)
            if isinstance(skills, list):
                _add_list(doc, [str(item) for item in skills])
            else:
                doc.add_paragraph(_safe_text(skills))

    resume = pipeline.get("resume", {}) or {}
    if resume:
        summary = resume.get("summary") or resume.get("professional_summary")
        if summary:
            doc.add_heading("Professional Summary", level=1)
            doc.add_paragraph(_safe_text(summary))

        experience = resume.get("experience_bullets") or resume.get("experience")
        if experience:
            doc.add_heading("Experience", level=1)
            if isinstance(experience, list):
                _add_list(doc, [str(item) for item in experience])
            else:
                doc.add_paragraph(_safe_text(experience))

        skills = resume.get("skills")
        if skills:
            doc.add_heading("Skills", level=1)
            if isinstance(skills, list):
                _add_list(doc, [str(item) for item in skills])
            else:
                doc.add_paragraph(_safe_text(skills))

        recommendations = resume.get("recommendations")
        if recommendations:
            doc.add_heading("Recommendations", level=1)
            if isinstance(recommendations, list):
                _add_list(doc, [str(item) for item in recommendations])
            else:
                doc.add_paragraph(_safe_text(recommendations))

    review = pipeline.get("review", {}) or {}
    if review:
        doc.add_heading("Review Notes", level=1)
        passed = review.get("passed")
        if passed is not None:
            doc.add_paragraph(f"Passed: {_safe_text(passed)}")
        issues = review.get("issues")
        if issues:
            doc.add_heading("Issues", level=2)
            if isinstance(issues, list):
                _add_list(doc, [str(item) for item in issues])
            else:
                doc.add_paragraph(_safe_text(issues))
        suggestions = review.get("suggestions")
        if suggestions:
            doc.add_heading("Suggestions", level=2)
            if isinstance(suggestions, list):
                _add_list(doc, [str(item) for item in suggestions])
            else:
                doc.add_paragraph(_safe_text(suggestions))

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
